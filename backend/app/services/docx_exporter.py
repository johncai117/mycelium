"""
DOCX exporter for EMA PASS / Non-Interventional Study Protocol format.

Formatting conventions derived from EMA-approved PASS protocols:
  - Arial 11pt body text, Arial 12pt bold section headings
  - Numbered headings matching EMA PASS template structure
  - Running header: protocol number + version (right-aligned), confidential (left)
  - Page numbers in footer (centred)
  - 1.15 line spacing throughout
  - Regulatory cover-page information table
  - Page break before each major (level-1) section
  - Table of Contents placeholder after cover block
  - 2.5 cm left / 2.0 cm right / 2.5 cm top / 2.0 cm bottom margins
"""

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.protocol import Protocol

# ---------------------------------------------------------------------------
# Section ordering and display titles (EMA PASS template numbering)
# ---------------------------------------------------------------------------

SECTION_ORDER = [
    "background",
    "objectives",
    "study_design",
    "study_setting",
    "cohort_definition",
    "variables",
    "study_size",
    "data_analysis",
    "limitations",
    "ethics",
]

# The cover page uses sections 1-2 for TOC + abbreviations; body starts at 3
SECTION_NUMBERS = {
    "background": 3,
    "objectives": 4,
    "study_design": 5,
    "study_setting": 6,
    "cohort_definition": 7,
    "variables": 8,
    "study_size": 9,
    "data_analysis": 10,
    "limitations": 11,
    "ethics": 12,
}

SECTION_TITLES = {
    "background": "Rationale and Background",
    "objectives": "Research Question and Objectives",
    "study_design": "Study Design",
    "study_setting": "Study Setting and Data Source",
    "cohort_definition": "Study Population and Cohort Definition",
    "variables": "Variables: Exposures, Outcomes, and Covariates",
    "study_size": "Study Size",
    "data_analysis": "Data Analysis",
    "limitations": "Limitations of the Research Methods",
    "ethics": "Protection of Human Participants",
}

# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _set_cell_background(cell, hex_color: str) -> None:
    """Fill a table cell with a solid background colour (hex, no #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_page_number_field(paragraph) -> None:
    """Insert a { PAGE } field run into *paragraph*."""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run._r.append(instrText)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_end)


def _add_page_break(doc: Document) -> None:
    """Add a hard page break paragraph."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    # Remove the paragraph spacing that would otherwise leave a blank line
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _set_repeat_table_header(row) -> None:
    """Mark a table row as a repeating header row."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


# ---------------------------------------------------------------------------
# Font / paragraph helpers
# ---------------------------------------------------------------------------

FONT_NAME = "Arial"
BODY_FONT_PT = 11
HEADING1_PT = 12
HEADING2_PT = 11
LINE_SPACING = 1.15   # multiple


def _apply_body_style(paragraph, bold: bool = False, italic: bool = False,
                      font_size: int = BODY_FONT_PT,
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    """Apply standard EMA body formatting to an existing paragraph."""
    pf = paragraph.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    paragraph.alignment = align
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic


def _add_body_paragraph(doc: Document, text: str, bold: bool = False,
                        italic: bool = False,
                        align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_FONT_PT)
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = align
    return p


def _add_section_heading(doc: Document, number: int, title: str,
                         level: int = 1) -> None:
    """Add a numbered section heading styled to EMA PASS conventions."""
    if level == 1:
        text = f"{number}. {title.upper()}"
        font_size = HEADING1_PT
    else:
        text = f"{number}. {title}"
        font_size = HEADING2_PT

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(font_size)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


# ---------------------------------------------------------------------------
# Header / footer
# ---------------------------------------------------------------------------

def _add_running_header(doc: Document, protocol_id: str, version: str) -> None:
    """Add a two-column running header: CONFIDENTIAL (left) | ID + Version (right)."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True  # cover page has no header

    header = section.header
    header.is_linked_to_previous = False

    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0

    # Left text: CONFIDENTIAL
    run_left = p.add_run("CONFIDENTIAL")
    run_left.font.name = FONT_NAME
    run_left.font.size = Pt(9)
    run_left.font.italic = True
    run_left.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Tab stop in the middle of the header paragraph → right-aligned text via tab
    # Use a right-aligned tab at page width
    tab = OxmlElement("w:tab")
    run_left._r.append(tab)

    # Right text: Protocol ID + Version
    run_right = p.add_run(f"{protocol_id}   |   Version {version}")
    run_right.font.name = FONT_NAME
    run_right.font.size = Pt(9)
    run_right.font.italic = True
    run_right.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Set a right-tab stop at 16 cm
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_stop = OxmlElement("w:tab")
    tab_stop.set(qn("w:val"), "right")
    tab_stop.set(qn("w:pos"), "9072")   # 16 cm in twentieths of a point
    tabs.append(tab_stop)
    pPr.append(tabs)

    # Horizontal rule below header
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_number_footer(doc: Document) -> None:
    """Add centred page number footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)

    run_pre = p.add_run("Page ")
    run_pre.font.name = FONT_NAME
    run_pre.font.size = Pt(9)

    _add_page_number_field(p)

    run_post = p.add_run(f"   |   {datetime.utcnow().strftime('%d %B %Y')}")
    run_post.font.name = FONT_NAME
    run_post.font.size = Pt(9)
    run_post.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ---------------------------------------------------------------------------
# Cover page / regulatory header table
# ---------------------------------------------------------------------------

def _add_cover_page(doc: Document, protocol: Protocol) -> None:
    """Render the EMA-style cover block on the first page."""
    si = protocol.study_inputs

    # Title block ─────────────────────────────────────────────────────────────
    drug_title = si.drug_name or "Investigational Medicinal Product"
    indication = si.indication or ""
    study_type_display = (si.study_type or "observational").replace("_", " ").title()

    p_drug = doc.add_paragraph()
    r = p_drug.add_run(drug_title.upper())
    r.font.name = FONT_NAME
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)   # dark navy
    p_drug.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_drug.paragraph_format.space_after = Pt(4)
    p_drug.paragraph_format.space_before = Pt(24)

    p_sub = doc.add_paragraph()
    r2 = p_sub.add_run(
        f"NON-INTERVENTIONAL STUDY PROTOCOL\n{protocol.study_id}"
    )
    r2.font.name = FONT_NAME
    r2.font.size = Pt(12)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(4)

    p_ver = doc.add_paragraph()
    r3 = p_ver.add_run(
        f"Version {protocol.version}.0,  "
        f"{datetime.utcnow().strftime('%d %B %Y')}"
    )
    r3.font.name = FONT_NAME
    r3.font.size = Pt(11)
    r3.font.italic = True
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ver.paragraph_format.space_after = Pt(20)

    # ── Regulatory information table ─────────────────────────────────────────
    long_title = getattr(si, "study_title", None) or (
        f"An Observational Study to Evaluate the Safety of {drug_title}"
        + (f" in {indication}" if indication else "")
    )

    sponsor_val = si.sponsor or "Not specified"
    reg_context = (si.regulatory_context or "PASS").replace("_", " ").upper()
    data_src = si.data_source or "Not specified"
    geography = si.geography or "Not specified"
    follow_up = (
        f"{getattr(si, 'follow_up_years', None)} year(s)"
        if getattr(si, "follow_up_years", None)
        else "To be determined"
    )
    primary_outcome = si.primary_outcome or "To be specified in protocol"

    rows = [
        ("Title", long_title),
        ("Protocol Number", protocol.study_id),
        ("Protocol Version Identifier", str(protocol.version)),
        ("Date", datetime.utcnow().strftime("%d %B %Y")),
        ("EU Post Authorisation Study (PAS) Register Number",
         "To be registered before the start of data collection"),
        ("Active Substance (INN)",
         si.drug_inn or si.drug_name),
        ("Medicinal Product", drug_title),
        ("Indication", indication or "Not specified"),
        ("Study Type", study_type_display),
        ("Primary Research Question / Objective", primary_outcome),
        ("Data Source(s)", data_src),
        ("Geography", geography),
        ("Follow-up Duration", follow_up),
        ("Marketing Authorisation Holder", sponsor_val),
        ("Regulatory Context", reg_context),
        ("Joint PASS", "No"),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row_idx, (label, value) in enumerate(rows):
        row = table.rows[row_idx]
        row.cells[0].width = Cm(6.5)
        row.cells[1].width = Cm(10.5)

        c0, c1 = row.cells[0], row.cells[1]
        _set_cell_background(c0, "1F3964")   # dark navy header cell

        # Label
        p0 = c0.paragraphs[0]
        p0.clear()
        r0 = p0.add_run(label)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(9)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)

        # Value
        p1 = c1.paragraphs[0]
        p1.clear()
        r1 = p1.add_run(str(value) if value is not None else "Not specified")
        r1.font.name = FONT_NAME
        r1.font.size = Pt(9)
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()   # spacing after table


# ---------------------------------------------------------------------------
# Table of Contents placeholder
# ---------------------------------------------------------------------------

def _add_toc_placeholder(doc: Document) -> None:
    _add_page_break(doc)

    p_title = doc.add_paragraph()
    r = p_title.add_run("1. TABLE OF CONTENTS")
    r.font.name = FONT_NAME
    r.font.size = Pt(HEADING1_PT)
    r.font.bold = True
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)

    p_note = doc.add_paragraph()
    r2 = p_note.add_run(
        "[Table of Contents — update field after final editing to populate entries]"
    )
    r2.font.name = FONT_NAME
    r2.font.size = Pt(BODY_FONT_PT)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p_note.paragraph_format.line_spacing = LINE_SPACING
    p_note.paragraph_format.space_after = Pt(4)

    # List the sections in the TOC placeholder
    for key in SECTION_ORDER:
        num = SECTION_NUMBERS[key]
        title = SECTION_TITLES[key]
        p_entry = doc.add_paragraph()
        r_entry = p_entry.add_run(f"{num}. {title}")
        r_entry.font.name = FONT_NAME
        r_entry.font.size = Pt(BODY_FONT_PT)
        p_entry.paragraph_format.left_indent = Cm(0.5)
        p_entry.paragraph_format.space_before = Pt(1)
        p_entry.paragraph_format.space_after = Pt(1)
        p_entry.paragraph_format.line_spacing = LINE_SPACING

    p_app = doc.add_paragraph()
    r_app = p_app.add_run("13. Appendix A: Code Sets")
    r_app.font.name = FONT_NAME
    r_app.font.size = Pt(BODY_FONT_PT)
    p_app.paragraph_format.left_indent = Cm(0.5)
    p_app.paragraph_format.space_before = Pt(1)
    p_app.paragraph_format.space_after = Pt(1)
    p_app.paragraph_format.line_spacing = LINE_SPACING


# ---------------------------------------------------------------------------
# Abbreviations section
# ---------------------------------------------------------------------------

def _add_abbreviations(doc: Document) -> None:
    _add_page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run("2. LIST OF ABBREVIATIONS")
    r.font.name = FONT_NAME
    r.font.size = Pt(HEADING1_PT)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    abbreviations = [
        ("AE", "Adverse event"),
        ("CI", "Confidence interval"),
        ("EMA", "European Medicines Agency"),
        ("EU", "European Union"),
        ("GVP", "Good Pharmacovigilance Practice"),
        ("HR", "Hazard ratio"),
        ("ICD", "International Classification of Diseases"),
        ("INN", "International Nonproprietary Name"),
        ("IRB", "Institutional Review Board / Ethics Committee"),
        ("MAH", "Marketing Authorisation Holder"),
        ("MedDRA", "Medical Dictionary for Regulatory Activities"),
        ("NIS", "Non-interventional study"),
        ("OR", "Odds ratio"),
        ("PAS", "Post-authorisation study"),
        ("PASS", "Post-authorisation safety study"),
        ("RMP", "Risk management plan"),
        ("RR", "Relative risk / rate ratio"),
        ("SAE", "Serious adverse event"),
        ("SAP", "Statistical analysis plan"),
        ("SmPC", "Summary of Product Characteristics"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0]
    _set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, ["Abbreviation", "Definition"]):
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.name = FONT_NAME
        r.font.size = Pt(BODY_FONT_PT)
        r.font.bold = True
        _set_cell_background(cell, "1F3964")
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for abbr, defn in abbreviations:
        row = table.add_row()
        for cell, text in zip(row.cells, [abbr, defn]):
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(text)
            r.font.name = FONT_NAME
            r.font.size = Pt(BODY_FONT_PT)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Code sets appendix
# ---------------------------------------------------------------------------

def _add_code_table(doc: Document, heading: str, items: list,
                    source_default: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(heading)
    r.font.name = FONT_NAME
    r.font.size = Pt(HEADING2_PT)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0]
    _set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, ["Code", "Description", "Source"]):
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.name = FONT_NAME
        r.font.size = Pt(BODY_FONT_PT)
        r.font.bold = True
        _set_cell_background(cell, "1F3964")
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for item in items:
        row = t.add_row()
        vals = [
            item.get("code", ""),
            item.get("description", ""),
            item.get("source", source_default),
        ]
        for cell, val in zip(row.cells, vals):
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(str(val))
            r.font.name = FONT_NAME
            r.font.size = Pt(BODY_FONT_PT)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate_docx(protocol: Protocol) -> bytes:
    doc = Document()

    # ── Page margins (EMA standard: 2.5 left, 2.0 right, 2.5 top, 2.0 bottom) ──
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.0)
        sec.header_distance = Cm(1.2)
        sec.footer_distance = Cm(1.2)

    # ── Running header + page-number footer ──────────────────────────────────
    _add_running_header(doc, protocol.study_id, str(protocol.version))
    _add_page_number_footer(doc)

    # ── Cover page ────────────────────────────────────────────────────────────
    _add_cover_page(doc, protocol)

    # ── Table of Contents placeholder ─────────────────────────────────────────
    _add_toc_placeholder(doc)

    # ── Abbreviations ─────────────────────────────────────────────────────────
    _add_abbreviations(doc)

    # ── Protocol sections ─────────────────────────────────────────────────────
    for section_key in SECTION_ORDER:
        _add_page_break(doc)

        num = SECTION_NUMBERS[section_key]
        title = SECTION_TITLES[section_key]
        _add_section_heading(doc, num, title, level=1)

        section_obj = protocol.sections.get(section_key)
        if section_obj and section_obj.content:
            for para_text in section_obj.content.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                _add_body_paragraph(doc, para_text)
        else:
            p = doc.add_paragraph()
            r = p.add_run("[Section not yet generated]")
            r.font.name = FONT_NAME
            r.font.size = Pt(BODY_FONT_PT)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.space_after = Pt(6)

        # Confidence / AI flag footnote
        if section_obj:
            if section_obj.ai_generated:
                p_flag = doc.add_paragraph()
                confidence_label = section_obj.confidence.capitalize()
                r_flag = p_flag.add_run(
                    f"[AI-generated content — confidence: {confidence_label}]"
                )
                r_flag.font.name = FONT_NAME
                r_flag.font.size = Pt(8)
                r_flag.font.italic = True
                r_flag.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
                p_flag.paragraph_format.space_before = Pt(0)
                p_flag.paragraph_format.space_after = Pt(2)

    # ── Appendix A: Code Sets ─────────────────────────────────────────────────
    _add_page_break(doc)
    p_app = doc.add_paragraph()
    r_app = p_app.add_run("13. APPENDIX A: CODE SETS")
    r_app.font.name = FONT_NAME
    r_app.font.size = Pt(HEADING1_PT)
    r_app.font.bold = True
    p_app.paragraph_format.space_before = Pt(12)
    p_app.paragraph_format.space_after = Pt(6)

    cs = protocol.code_sets
    has_codes = False

    if cs.icd10:
        has_codes = True
        _add_code_table(doc, "A.1  ICD-10 Diagnosis Codes",
                        cs.icd10, "ICD-10-CM")

    if cs.ndc:
        has_codes = True
        _add_code_table(doc, "A.2  NDC / Drug Codes", cs.ndc, "NDC")

    if cs.cpt:
        has_codes = True
        _add_code_table(doc, "A.3  CPT Procedure Codes", cs.cpt, "CPT")

    if not has_codes:
        p_none = doc.add_paragraph()
        r_none = p_none.add_run("No code sets have been specified for this protocol.")
        r_none.font.name = FONT_NAME
        r_none.font.size = Pt(BODY_FONT_PT)
        r_none.font.italic = True
        p_none.paragraph_format.line_spacing = LINE_SPACING
        p_none.paragraph_format.space_after = Pt(6)

    # ── Serialize ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
